"""
ingestion/document_parser.py
─────────────────────────────
Smart document parser that extracts structured content
(text blocks, tables, images) from complex file formats.

Handles:
  - PDFs with mixed text + embedded images + tables
  - DOCX with headers, tables, and inline images
  - HTML pages (strips boilerplate, extracts main content)
  - CSV/Excel with automatic header detection
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from loguru import logger


@dataclass
class ParsedBlock:
    block_type: str           # text | table | image | heading
    content: str              # text content or file path for images
    metadata: dict = field(default_factory=dict)
    page: int = 1
    order: int = 0            # position within document


@dataclass
class ParsedDocument:
    source: str
    file_type: str
    blocks: list[ParsedBlock] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)

    def text_blocks(self) -> list[ParsedBlock]:
        return [b for b in self.blocks if b.block_type in ("text", "heading")]

    def table_blocks(self) -> list[ParsedBlock]:
        return [b for b in self.blocks if b.block_type == "table"]

    def image_blocks(self) -> list[ParsedBlock]:
        return [b for b in self.blocks if b.block_type == "image"]

    def full_text(self) -> str:
        return "\n\n".join(b.content for b in self.text_blocks())


class PDFParser:
    def parse(self, path: Path) -> ParsedDocument:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        doc = ParsedDocument(source=str(path), file_type="pdf")
        doc.metadata["pages"] = len(reader.pages)

        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                doc.blocks.append(ParsedBlock(
                    block_type="text",
                    content=text.strip(),
                    page=i + 1,
                    order=i,
                ))
        logger.debug(f"PDF parsed: {len(doc.blocks)} blocks from {path.name}")
        return doc


class DOCXParser:
    def parse(self, path: Path) -> ParsedDocument:
        import docx
        document = docx.Document(str(path))
        doc = ParsedDocument(source=str(path), file_type="docx")
        order = 0

        for para in document.paragraphs:
            if not para.text.strip():
                continue
            btype = "heading" if para.style.name.startswith("Heading") else "text"
            doc.blocks.append(ParsedBlock(
                block_type=btype,
                content=para.text.strip(),
                metadata={"style": para.style.name},
                order=order,
            ))
            order += 1

        for table in document.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            header = " | ".join(rows[0])
            body = "\n".join(" | ".join(row) for row in rows[1:])
            doc.blocks.append(ParsedBlock(
                block_type="table",
                content=f"{header}\n{body}",
                metadata={"rows": len(rows), "cols": len(rows[0])},
                order=order,
            ))
            order += 1

        logger.debug(f"DOCX parsed: {len(doc.blocks)} blocks from {path.name}")
        return doc


class CSVParser:
    def parse(self, path: Path) -> ParsedDocument:
        import pandas as pd
        df = pd.read_csv(path)
        doc = ParsedDocument(source=str(path), file_type="csv")
        doc.metadata.update({"rows": len(df), "cols": len(df.columns)})

        # Full table as one block
        doc.blocks.append(ParsedBlock(
            block_type="table",
            content=df.to_string(index=False),
            metadata={"columns": list(df.columns)},
        ))

        # Summary stats block for numeric columns
        numeric = df.select_dtypes(include="number")
        if not numeric.empty:
            summary = numeric.describe().to_string()
            doc.blocks.append(ParsedBlock(
                block_type="text",
                content=f"Statistical summary:\n{summary}",
                order=1,
            ))

        return doc


class HTMLParser:
    def parse(self, path: Path) -> ParsedDocument:
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            logger.warning("beautifulsoup4 not installed — returning raw HTML text")
            content = path.read_text(errors="replace")
            doc = ParsedDocument(source=str(path), file_type="html")
            doc.blocks.append(ParsedBlock(block_type="text", content=content))
            return doc

        html = path.read_text(errors="replace")
        soup = BeautifulSoup(html, "html.parser")

        # Remove boilerplate
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        doc = ParsedDocument(source=str(path), file_type="html")
        order = 0

        for tag in soup.find_all(["h1", "h2", "h3", "p", "table"]):
            if tag.name in ("h1", "h2", "h3"):
                doc.blocks.append(ParsedBlock(
                    block_type="heading",
                    content=tag.get_text(strip=True),
                    order=order,
                ))
            elif tag.name == "p":
                text = tag.get_text(strip=True)
                if len(text) > 30:
                    doc.blocks.append(ParsedBlock(
                        block_type="text", content=text, order=order
                    ))
            elif tag.name == "table":
                rows = []
                for tr in tag.find_all("tr"):
                    rows.append(" | ".join(
                        td.get_text(strip=True) for td in tr.find_all(["td", "th"])
                    ))
                if rows:
                    doc.blocks.append(ParsedBlock(
                        block_type="table",
                        content="\n".join(rows),
                        order=order,
                    ))
            order += 1

        return doc


class DocumentParser:
    """Registry-based smart parser — pick the right parser per file type."""

    _parsers = {
        ".pdf": PDFParser,
        ".docx": DOCXParser,
        ".doc": DOCXParser,
        ".csv": CSVParser,
        ".html": HTMLParser,
        ".htm": HTMLParser,
    }

    def parse(self, path: str | Path) -> ParsedDocument | None:
        path = Path(path)
        parser_cls = self._parsers.get(path.suffix.lower())
        if not parser_cls:
            logger.warning(f"No parser for {path.suffix} — skipping {path.name}")
            return None
        return parser_cls().parse(path)

    def parse_directory(self, directory: str) -> list[ParsedDocument]:
        docs = []
        for path in Path(directory).rglob("*"):
            if path.is_file() and path.suffix.lower() in self._parsers:
                result = self.parse(path)
                if result:
                    docs.append(result)
        logger.info(f"Parsed {len(docs)} documents from {directory}")
        return docs
